# backend/routes/synopsis.py
"""
Camp Synopsis routes — weeks, camps, entries, photos, and doc generation.
Public read/write for entries and photos; ICC admin required for write on weeks/camps.
"""

import io
import logging
import os as _os
import re
import urllib.parse
import urllib.request
import uuid
import zipfile
import datetime as _dt
from datetime import datetime
from html import escape as _html_escape
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement

from fastapi import APIRouter, Depends, File, Header, HTTPException, Query, UploadFile, status
from fastapi.responses import StreamingResponse

import firebase_admin
from firebase_admin import auth

from schemas.synopsis_schema import (
    CampCreate,
    CampUpdate,
    CampGroupInput,
    EnhanceTextRequest,
    EntrySaveRequest,
    EntryDayInput,
    FoodDayData,
    FoodUpdateRequest,
    SubCampInput,
    SynopsisCampFields,
    SynopsisEntryFields,
    SynopsisWeekFields,
    WeekCreate,
    WeekSetupRequest,
    WeekUpdate,
    ENTRY_STATUS_SAVED,
    ICC_ADMIN_DOMAIN,
    PHOTO_MIN,
    PHOTO_MAX,
    VALID_DAYS,
)
from services.firebase_service import FirebaseService
from utils.llm_handler import call_openai

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/synopsis", tags=["synopsis"])
firebase = FirebaseService()

ENHANCE_SYSTEM_PROMPT = (
    "You help teachers write clear, friendly camp activity summaries for parents at a K-8 enrichment program. "
    "Paraphrase the teacher's note in plain, warm language. Rules:\n"
    "- Only include what the teacher wrote. Do not add, invent, or elaborate on anything.\n"
    "- Keep it concise: 1 to 2 short paragraphs.\n"
    "- Do not use em dashes (—). Use commas or short sentences instead.\n"
    "- Add 1 to 2 relevant emojis placed naturally within the text.\n"
    "- Bold 1 to 2 key phrases using markdown bold (**like this**).\n"
    "- Keep a warm, first-person plural voice (e.g. 'we built...', 'our campers...').\n"
    "- Return only the paraphrased text, no preamble, no quotes around it."
)

LIGHT_CORRECT_PROMPT = (
    "You are a copy editor for a K-8 enrichment camp newsletter. "
    "A teacher has written a brief activity description for parents. "
    "Lightly polish the text so it reads clearly and professionally, while keeping the teacher's facts, voice, and all details fully intact. Rules:\n"
    "- Fix all spelling, grammar, and punctuation errors.\n"
    "- Improve sentence flow and readability where helpful (e.g. break run-on sentences, smooth awkward phrasing) "
    "without paraphrasing or restructuring the overall content.\n"
    "- Do NOT add or remove any activity, fact, or detail the teacher wrote.\n"
    "- Keep the teacher's warm, first-person voice.\n"
    "- Add 1 to 2 relevant emojis placed naturally within the text.\n"
    "- Return only the polished text. No preamble, no explanation, no quotes around it."
)

DAY_LABELS = {'mon': 'Monday', 'tue': 'Tuesday', 'wed': 'Wednesday', 'thu': 'Thursday', 'fri': 'Friday'}
DAY_ORDER = list(VALID_DAYS)

PARSE_FOOD_PROMPT = (
    "You extract a weekly camp food menu from an image. "
    "Return JSON with exactly this structure (leave fields as empty strings if not found):\n"
    '{"mon": {"morning_snack": "", "lunch": "", "afternoon_snack": ""}, '
    '"tue": {"morning_snack": "", "lunch": "", "afternoon_snack": ""}, '
    '"wed": {"morning_snack": "", "lunch": "", "afternoon_snack": ""}, '
    '"thu": {"morning_snack": "", "lunch": "", "afternoon_snack": ""}, '
    '"fri": {"morning_snack": "", "lunch": "", "afternoon_snack": ""}}\n\n'
    "Fill in every field you can read. Use short, plain descriptions (e.g. 'Fresh fruit and crackers'). "
    "Leave a field as an empty string if the menu does not mention it."
)

PARSE_FLYER_PROMPT = (
    "You extract camp schedule information from a flyer, image, or spreadsheet. "
    "Return JSON with exactly this structure (leave fields as empty strings if not found):\n"
    '{"week_name": "string", "start_date": "YYYY-MM-DD or empty", "end_date": "YYYY-MM-DD or empty", '
    '"camp_groups": [{"group_name": "string", "age_group": "e.g. Ages 5-8 or Grades 1-3 or empty", "camps": ['
    '{"camp_name": "string", "time_start": "e.g. 9:00 AM", "time_end": "e.g. 12:00 PM"}]}]}\n\n'
    "IMPORTANT field definitions:\n"
    "- group_name: The FULL camp program name as written on the flyer (e.g. 'Brick Moto Mighty & Sports Camp', "
    "'Eco Art & Music Camp'). This is NOT an age range or grade level.\n"
    "- age_group: Any age range, grade level, or demographic label associated with the camp group "
    "(e.g. 'Ages 5-8', 'Grades 1-3', 'K-2', 'Junior', 'Senior'). Leave empty if not mentioned.\n"
    "- camp_name: Each individual sub-camp within the group, parsed from the group_name by splitting on '&', "
    "'and', or commas (e.g. group 'Brick Moto Mighty & Sports Camp' → camps 'Brick Moto Mighty' and 'Sports'). "
    "Each sub-camp has its own time slot.\n"
    "- Do NOT put age ranges, grade levels, or demographic info in group_name. Put them in age_group instead.\n"
    "- If a camp program has no '&' or 'and' in the name, it has one sub-camp entry with the same name (minus 'Camp')."
)


# ── Auth dependency ───────────────────────────────────────────────────────────

async def verify_icc_admin(authorization: Optional[str] = Header(default=None)) -> dict:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="ICC admin access required")
    try:
        token = authorization.split("Bearer ")[1]
        decoded = auth.verify_id_token(token)
        email = decoded.get("email", "")
        if not email.endswith(f"@{ICC_ADMIN_DOMAIN}"):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="ICC admin access required")
        return decoded
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="ICC admin access required")



# ── Weeks ─────────────────────────────────────────────────────────────────────

@router.get("/weeks")
async def list_weeks():
    weeks = await firebase.get_all_synopsis_weeks()
    return {"weeks": weeks}


@router.get("/weeks/active")
async def get_active_week():
    week = await firebase.get_active_synopsis_week()
    if not week:
        return {"week": None}
    return {"week": week}


@router.get("/weeks/visible")
async def list_visible_weeks():
    """Weeks the admin has flagged as visible to teachers — used to populate the teacher week picker.

    The active week is always included even if it isn't explicitly flagged
    visible, so weeks created before this feature (or where the admin forgot
    to tick the box) still reach teachers.
    """
    weeks = await firebase.get_visible_synopsis_weeks()
    active = await firebase.get_active_synopsis_week()
    if active:
        active_id = active.get(SynopsisWeekFields.WEEK_ID) or active.get('id')
        already_included = any((w.get(SynopsisWeekFields.WEEK_ID) or w.get('id')) == active_id for w in weeks)
        if not already_included:
            weeks.insert(0, active)
    return {"weeks": weeks}


@router.post("/weeks", status_code=status.HTTP_201_CREATED)
async def create_week(body: WeekCreate, admin: dict = Depends(verify_icc_admin)):
    now = datetime.utcnow().isoformat()
    if body.is_active:
        await firebase.deactivate_all_synopsis_weeks()
    data = {
        SynopsisWeekFields.LABEL: body.label,
        SynopsisWeekFields.START_DATE: body.start_date,
        SynopsisWeekFields.END_DATE: body.end_date,
        SynopsisWeekFields.IS_ACTIVE: body.is_active,
        SynopsisWeekFields.IS_VISIBLE: body.is_visible,
        SynopsisWeekFields.DRIVE_LINK: body.drive_link or "",
        SynopsisWeekFields.CREATED_BY: admin.get("uid", ""),
        SynopsisWeekFields.CREATED_AT: now,
    }
    week_id = await firebase.create_synopsis_week(data)
    return {"week_id": week_id}


@router.patch("/weeks/{week_id}")
async def update_week(week_id: str, body: WeekUpdate, admin: dict = Depends(verify_icc_admin)):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(400, "No fields to update")
    if updates.get(SynopsisWeekFields.IS_ACTIVE):
        await firebase.deactivate_all_synopsis_weeks()
    ok = await firebase.update_synopsis_week(week_id, updates)
    if not ok:
        raise HTTPException(404, "Week not found")
    return {"success": True}


@router.delete("/weeks/{week_id}")
async def delete_week(week_id: str, admin: dict = Depends(verify_icc_admin)):
    ok = await firebase.delete_synopsis_week(week_id)
    if not ok:
        raise HTTPException(404, "Week not found")
    return {"success": True}


@router.post("/weeks/setup", status_code=status.HTTP_201_CREATED)
async def setup_week(body: WeekSetupRequest, admin: dict = Depends(verify_icc_admin)):
    """Create a week and all its camps in a single call."""
    now = datetime.utcnow().isoformat()
    if body.is_active:
        await firebase.deactivate_all_synopsis_weeks()

    week_data = {
        SynopsisWeekFields.LABEL: body.label,
        SynopsisWeekFields.START_DATE: body.start_date,
        SynopsisWeekFields.END_DATE: body.end_date,
        SynopsisWeekFields.IS_ACTIVE: body.is_active,
        SynopsisWeekFields.IS_VISIBLE: body.is_visible,
        SynopsisWeekFields.DRIVE_LINK: body.drive_link or "",
        SynopsisWeekFields.CREATED_BY: admin.get("uid", ""),
        SynopsisWeekFields.CREATED_AT: now,
    }
    week_id = await firebase.create_synopsis_week(week_data)

    camp_ids = []
    for group in body.camp_groups:
        for camp in group.camps:
            camp_data = {
                SynopsisCampFields.WEEK_ID: week_id,
                SynopsisCampFields.GROUP_NAME: group.group_name,
                SynopsisCampFields.AGE_GROUP: group.age_group or "",
                SynopsisCampFields.CAMP_NAME: camp.camp_name,
                SynopsisCampFields.TEACHER_NAME: camp.teacher_name or "",
                SynopsisCampFields.TIME_START: camp.time_start or "",
                SynopsisCampFields.TIME_END: camp.time_end or "",
                SynopsisCampFields.CREATED_AT: now,
            }
            camp_id = await firebase.create_synopsis_camp(camp_data)
            camp_ids.append(camp_id)

    return {"week_id": week_id, "camp_ids": camp_ids}


@router.post("/parse-flyer")
async def parse_flyer(
    file: UploadFile = File(...),
    admin: dict = Depends(verify_icc_admin),
):
    """Upload a camp flyer image or Excel file; returns structured week/camp data."""
    import base64, os, tempfile

    content_type = (file.content_type or "").lower()
    data = await file.read()

    if content_type.startswith("image/"):
        b64 = base64.b64encode(data).decode()
        data_uri = f"data:{content_type};base64,{b64}"
        result = call_openai(
            prompt="Extract the camp schedule information from this image.",
            system_message=PARSE_FLYER_PROMPT,
            json_mode=True,
            images=[data_uri],
        )
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "text/csv",
        "application/octet-stream",
    ):
        suffix = ".xlsx" if "spreadsheet" in content_type or "octet" in content_type else ".xls"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            from services.file_parser import extract_text_from_file
            text = extract_text_from_file(tmp_path, suffix)
        finally:
            os.remove(tmp_path)
        result = call_openai(
            prompt=f"Extract the camp schedule information from this document:\n\n{text}",
            system_message=PARSE_FLYER_PROMPT,
            json_mode=True,
        )
    else:
        raise HTTPException(400, "Please upload an image (JPG, PNG) or Excel file (.xlsx)")

    # Ensure the response has the expected shape
    result.setdefault("week_name", "")
    result.setdefault("start_date", "")
    result.setdefault("end_date", "")
    result.setdefault("camp_groups", [])
    return result


@router.post("/parse-food")
async def parse_food_image(
    file: UploadFile = File(...),
    admin: dict = Depends(verify_icc_admin),
):
    """Upload a menu photo; returns structured food data keyed by day."""
    import base64
    content_type = (file.content_type or "").lower()
    if not content_type.startswith("image/"):
        raise HTTPException(400, "Please upload an image (JPG, PNG, etc.)")
    data = await file.read()
    b64      = base64.b64encode(data).decode()
    data_uri = f"data:{content_type};base64,{b64}"
    result = call_openai(
        prompt="Extract the camp food menu from this image.",
        system_message=PARSE_FOOD_PROMPT,
        json_mode=True,
        images=[data_uri],
    )
    for day in ("mon", "tue", "wed", "thu", "fri"):
        result.setdefault(day, {"morning_snack": "", "lunch": "", "afternoon_snack": ""})
    return result


# ── Camps ─────────────────────────────────────────────────────────────────────

@router.get("/weeks/{week_id}/camps")
async def list_camps(week_id: str):
    camps = await firebase.get_synopsis_camps_for_week(week_id)
    return {"camps": camps}


@router.post("/camps", status_code=status.HTTP_201_CREATED)
async def create_camp(body: CampCreate, admin: dict = Depends(verify_icc_admin)):
    data = {
        SynopsisCampFields.WEEK_ID: body.week_id,
        SynopsisCampFields.GROUP_NAME: body.group_name or "",
        SynopsisCampFields.AGE_GROUP: body.age_group or "",
        SynopsisCampFields.CAMP_NAME: body.camp_name,
        SynopsisCampFields.TEACHER_NAME: body.teacher_name or "",
        SynopsisCampFields.TIME_START: body.time_start or "",
        SynopsisCampFields.TIME_END: body.time_end or "",
        SynopsisCampFields.CREATED_AT: datetime.utcnow().isoformat(),
    }
    camp_id = await firebase.create_synopsis_camp(data)
    return {"camp_id": camp_id}


@router.patch("/camps/{camp_id}")
async def update_camp(camp_id: str, body: CampUpdate, week_id: str = Query(...), admin: dict = Depends(verify_icc_admin)):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(400, "No fields to update")
    ok = await firebase.update_synopsis_camp(week_id, camp_id, updates)
    if not ok:
        raise HTTPException(404, "Camp not found")
    return {"success": True}


@router.delete("/camps/{camp_id}")
async def delete_camp(camp_id: str, week_id: str = Query(...), admin: dict = Depends(verify_icc_admin)):
    ok = await firebase.delete_synopsis_camp(week_id, camp_id)
    if not ok:
        raise HTTPException(404, "Camp not found")
    return {"success": True}


# ── Entries ───────────────────────────────────────────────────────────────────

@router.get("/camps/{camp_id}/entries")
async def get_entries(camp_id: str, week_id: Optional[str] = Query(None)):
    # week_id lets callers target any week's camps (needed now that teachers can
    # add synopses to any admin-visible week, not just the single "active" one).
    if not week_id:
        active_week = await firebase.get_active_synopsis_week()
        if not active_week:
            return {"entries": []}
        week_id = active_week.get(SynopsisWeekFields.WEEK_ID) or active_week.get('id')
    entries = await firebase.get_synopsis_entries_for_camp_in_week(week_id, camp_id)
    return {"entries": entries}


@router.get("/entries/{entry_id}")
async def get_entry(entry_id: str, week_id: Optional[str] = Query(None), camp_id: Optional[str] = Query(None)):
    if week_id and camp_id:
        doc = firebase._entries_col(week_id, camp_id).document(entry_id).get()
        entry = ({'id': doc.id, **doc.to_dict()} if doc.exists else None)
    else:
        entry = await firebase.get_synopsis_entry(entry_id)
    if not entry:
        raise HTTPException(404, "Entry not found")
    return entry


@router.post("/entries")
async def save_entries(body: EntrySaveRequest):
    # Teachers now pick which (admin-visible) week they're submitting for, so the
    # target week comes from the request instead of always being the active week.
    week_id = body.week_id

    # Verify the camp exists within the given week (direct path, no composite index needed)
    camp = await firebase.get_synopsis_camp_in_week(week_id, body.camp_id)
    if not camp:
        raise HTTPException(404, "Camp not found in the given week")

    saved_ids = []

    for entry in body.entries:
        day = entry.day.lower()
        if day not in VALID_DAYS:
            raise HTTPException(400, f"Invalid day '{day}'. Must be one of: {', '.join(VALID_DAYS)}")

        photo_count = len(entry.photo_urls)
        if photo_count > 0 and (photo_count < PHOTO_MIN or photo_count > PHOTO_MAX):
            raise HTTPException(
                400,
                f"Day '{day}': photos must be between {PHOTO_MIN} and {PHOTO_MAX} (got {photo_count})"
            )

        now = datetime.utcnow().isoformat()
        data = {
            SynopsisEntryFields.DAY_TITLE: entry.day_title,
            SynopsisEntryFields.RAW_TEXT: entry.raw_text,
            SynopsisEntryFields.PHOTO_URLS: entry.photo_urls,
            SynopsisEntryFields.STATUS: ENTRY_STATUS_SAVED,
            SynopsisEntryFields.UPDATED_AT: now,
        }

        entry_id = await firebase.upsert_synopsis_entry(body.camp_id, week_id, day, data)
        saved_ids.append(entry_id)

    return {"success": True, "entry_ids": saved_ids}


@router.post("/enhance")
async def enhance_text(body: EnhanceTextRequest):
    """Enhance a day's description with AI — teacher-triggered only."""
    if not body.raw_text.strip():
        raise HTTPException(400, "raw_text is required")
    result = call_openai(
        prompt=body.raw_text,
        system_message=ENHANCE_SYSTEM_PROMPT,
        json_mode=False,
        temperature=0.7,
    )
    return {"enhanced_text": result.get("response", "")}


# ── Photos ────────────────────────────────────────────────────────────────────

ALLOWED_IMAGE_TYPES = {"image/jpeg", "image/jpg", "image/png", "image/webp", "image/gif"}
MAX_PHOTO_BYTES = 10 * 1024 * 1024  # 10 MB


@router.post("/photos")
async def upload_photo(
    file: UploadFile = File(...),
    camp_id: str = Query(...),
    day: str = Query(...),
):
    day = day.lower()
    if day not in VALID_DAYS:
        raise HTTPException(400, f"Invalid day. Must be one of: {', '.join(VALID_DAYS)}")

    content_type = file.content_type or ""
    if content_type not in ALLOWED_IMAGE_TYPES:
        raise HTTPException(400, "Only JPEG, PNG, WebP, and GIF images are supported")

    data = await file.read()
    if len(data) > MAX_PHOTO_BYTES:
        raise HTTPException(400, "Image must be under 10 MB")

    ext = file.filename.rsplit(".", 1)[-1] if "." in file.filename else "jpg"
    photo_id = str(uuid.uuid4())
    storage_path = f"synopsis/{camp_id}/{day}/{photo_id}.{ext}"

    url = await firebase.upload_file(data, storage_path, content_type)
    return {"url": url}


@router.delete("/photos")
async def delete_photo(url: str = Query(...)):
    # Best-effort: we don't track photos in a separate collection so just return success.
    # The client removes the URL from its local state; on next save it won't be persisted.
    return {"success": True}


# ── Food ─────────────────────────────────────────────────────────────────────

@router.get("/weeks/{week_id}/food")
async def get_food(week_id: str):
    food = await firebase.get_synopsis_food(week_id)
    return {"food": food}


@router.patch("/weeks/{week_id}/food")
async def update_food(week_id: str, body: FoodUpdateRequest):
    data = {day: vals.model_dump() for day, vals in body.days.items()}
    await firebase.upsert_synopsis_food(week_id, data)
    return {"success": True}


# ── Doc generation ────────────────────────────────────────────────────────────

def _fetch_photo_bytes(url: str) -> bytes:
    """Download a photo for doc embedding.

    For Firebase Storage URLs we extract the blob path and use the Admin SDK
    directly (avoids public-URL auth issues on Cloud Run).  Falls back to an
    HTTP GET for any other URL.
    """
    parsed = urllib.parse.urlparse(url)
    if "firebasestorage.googleapis.com" in parsed.netloc:
        # URL shape: /v0/b/{bucket}/o/{encoded_path}?alt=media&token=...
        parts = parsed.path.split("/o/", 1)
        if len(parts) == 2:
            blob_path = urllib.parse.unquote(parts[1])
            try:
                return firebase.bucket.blob(blob_path).download_as_bytes(timeout=20)
            except Exception as exc:
                logger.warning(f"SDK download failed for {blob_path}: {exc}; falling back to HTTP")

    # HTTP fallback (also used for non-Firebase URLs)
    import requests as _requests
    resp = _requests.get(url, headers={"User-Agent": "EdCube/1.0"}, timeout=20)
    resp.raise_for_status()
    return resp.content


def _fix_exif_rotation(img_bytes: bytes) -> bytes:
    """Auto-rotate image based on EXIF orientation so it appears right-side up."""
    try:
        from PIL import Image as PILImage, ImageOps
        img = PILImage.open(io.BytesIO(img_bytes))
        img = ImageOps.exif_transpose(img)
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        out = io.BytesIO()
        img.save(out, format='JPEG', quality=85)
        out.seek(0)
        return out.read()
    except Exception:
        return img_bytes




def _add_md_text(para, text: str, size_pt: int = 12) -> None:
    """Add text to a Word paragraph, converting **bold** markdown to actual bold runs."""
    parts = re.split(r'\*\*', text or '')
    for i, part in enumerate(parts):
        if part:
            _run(para, part, bold=(i % 2 == 1), size_pt=size_pt)


def _light_correct(text: str) -> str:
    """Fix spelling/grammar/punctuation and add emojis without changing meaning."""
    if not text or not text.strip():
        return text
    try:
        result = call_openai(
            prompt=text,
            system_message=LIGHT_CORRECT_PROMPT,
            json_mode=False,
            temperature=0.1,
        )
        return result.get("response", text) or text
    except Exception:
        logger.warning("Light correction failed; using original teacher text.")
        return text


# Day ordinal labels used in the document (matches reference doc format)
DAY_ORDINAL = {'mon': 'Day 1', 'tue': 'Day 2', 'wed': 'Day 3', 'thu': 'Day 4', 'fri': 'Day 5'}

# Document colour constants (from reference DOCX)
COLOR_TITLE    = "1C4587"  # dark blue — title
COLOR_FOOD_BG  = "f9cb9c"  # orange — food day/snack rows
COLOR_GROUP_BG = "f6b26b"  # darker orange — group camp heading
COLOR_CAMP_BG  = "ffe599"  # light gold — sub-camp name label
COLOR_DAY_BG   = "dd7e6b"  # red-orange — day heading within camp


def _run(para, text: str, bold=False, italic=False, size_pt=12,
         color_hex: str = None, bg_hex: str = None):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.font.name = "Lora"
    r.font.size = Pt(size_pt)
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)
    if bg_hex:
        rPr = r._r.get_or_add_rPr()
        shd = OxmlElement('w:shd')
        shd.set(_qn('w:val'),   'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'),  bg_hex)
        rPr.append(shd)
    return r


def _para(doc, style='Normal', align=None, spc_b=None, spc_a=None, line=None, para_bg=None):
    p  = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if spc_b is not None:
        pf.space_before = Pt(spc_b)
    if spc_a is not None:
        pf.space_after  = Pt(spc_a)
    if line is not None:
        pf.line_spacing = line
    if para_bg:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(_qn('w:val'),   'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'),  para_bg)
        pPr.append(shd)
    return p


def _cell_bg(cell, hex6: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(_qn('w:val'),   'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'),  hex6.lstrip('#'))
    tcPr.append(shd)


def _para_border_bottom(para, color_hex: str = '1a1a1a', sz: str = '8'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(_qn('w:val'),   'single')
    btm.set(_qn('w:sz'),    sz)
    btm.set(_qn('w:space'), '1')
    btm.set(_qn('w:color'), color_hex.lstrip('#'))
    pBdr.append(btm)
    pPr.append(pBdr)


def _remove_table_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tcPr  = cell._tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(_qn('w:val'), 'none')
                tcBdr.append(b)
            tcPr.append(tcBdr)


_HYPERLINK_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'


def _add_hyperlink(para, url: str, display_text: str, size_pt: int = 10, bold: bool = True):
    """Insert a clickable hyperlink run into an existing paragraph."""
    rId = para.part.relate_to(url, _HYPERLINK_REL, is_external=True)
    hl  = OxmlElement('w:hyperlink')
    hl.set(_qn('r:id'), rId)
    r   = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(_qn('w:ascii'),    'Lora')
    rFonts.set(_qn('w:hAnsi'),   'Lora')
    rPr.append(rFonts)
    sz_el = OxmlElement('w:sz')
    sz_el.set(_qn('w:val'), str(size_pt * 2))
    rPr.append(sz_el)
    if bold:
        rPr.append(OxmlElement('w:b'))
    color_el = OxmlElement('w:color')
    color_el.set(_qn('w:val'), '0563C1')
    rPr.append(color_el)
    u_el = OxmlElement('w:u')
    u_el.set(_qn('w:val'), 'single')
    rPr.append(u_el)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    r.append(t)
    hl.append(r)
    para._p.append(hl)


# ── Shared doc-generation constants ──────────────────────────────────────────
_LOGO_PATH = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), 'assets', 'ICC_official_logo.png')
_DAY_NUM         = {'mon': '1', 'tue': '2', 'wed': '3', 'thu': '4', 'fri': '5'}
_DAY_C_HEX       = {'mon': 'B8E8A5', 'tue': 'A5C9E8', 'wed': 'E8D5A5', 'thu': 'E8A5A5', 'fri': 'C5B8E8'}
_DAY_LABEL_SHORT = {'mon': 'Mon',    'tue': 'Tue',    'wed': 'Wed',    'thu': 'Thu',    'fri': 'Fri'}
_DAY_EMOJI       = {'mon': '🌟',     'tue': '⭐',     'wed': '🌈',     'thu': '🎉',     'fri': '🎊'}


def _build_synopsis_doc(
    *,
    group_camps: list,
    doc_title: str,
    header_meta: str,
    drive_link: str,
    entries_by_camp: dict,
    food_data: dict,
    year: int,
) -> bytes:
    """Build one synopsis .docx for a single camp group and return raw bytes."""
    doc     = Document()
    section = doc.sections[0]
    section.top_margin      = Inches(1.1)
    section.bottom_margin   = Inches(1.1)
    section.left_margin     = Inches(0.85)
    section.right_margin    = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    # Header / footer
    for hf_p in (section.header.paragraphs[0], section.footer.paragraphs[0]):
        hf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hf_r = hf_p.add_run(f'ICC Summer Camps Program {year}')
        hf_r.font.size      = Pt(11)
        hf_r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Title block: camp name + meta | logo
    ttbl = doc.add_table(rows=1, cols=2)
    ttbl.autofit = False
    ttbl.columns[0].width = Inches(4.5)
    ttbl.columns[1].width = Inches(2.0)
    _remove_table_borders(ttbl)

    cl     = ttbl.cell(0, 0)
    p_name = cl.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(6)
    p_name.paragraph_format.space_after  = Pt(6)
    r_title = p_name.add_run(doc_title)
    r_title.bold           = True
    r_title.font.size      = Pt(22)
    r_title.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    p_meta = cl.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(4)
    p_meta.paragraph_format.space_after  = Pt(10)
    r_meta = p_meta.add_run(header_meta)
    r_meta.font.size      = Pt(12)
    r_meta.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    cr     = ttbl.cell(0, 1)
    p_logo = cr.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if _os.path.exists(_LOGO_PATH):
        try:
            p_logo.add_run().add_picture(_LOGO_PATH, width=Inches(1.9))
        except Exception:
            pass

    p_div = _para(doc, spc_b=6, spc_a=10)
    _para_border_bottom(p_div, '1a1a1a', sz='12')

    # Gallery bar with clickable hyperlink
    if drive_link:
        p_gal = _para(doc, spc_b=4, spc_a=8)
        pPr = p_gal._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(_qn('w:val'),   'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'),  'A5C9E8')
        pPr.append(shd)
        _run(p_gal, '📷  This week\'s photos — ', size_pt=10)
        _add_hyperlink(p_gal, drive_link, drive_link, size_pt=10, bold=True)

    # Camp sections
    for camp_idx, camp in enumerate(group_camps):
        camp_id      = camp.get(SynopsisCampFields.CAMP_ID)
        camp_name    = camp.get(SynopsisCampFields.CAMP_NAME, 'Camp')
        camp_entries = entries_by_camp.get(camp_id, {})

        p_camp = _para(doc, spc_b=16, spc_a=10, align=WD_ALIGN_PARAGRAPH.CENTER, para_bg=COLOR_CAMP_BG)
        if camp_idx > 0:
            p_camp.paragraph_format.page_break_before = True
        _run(p_camp, camp_name, bold=True, size_pt=23)
        _para_border_bottom(p_camp, '1a1a1a', sz='6')

        for day_key in DAY_ORDER:
            entry     = camp_entries.get(day_key)
            day_color = _DAY_C_HEX[day_key]
            day_num   = _DAY_NUM[day_key]
            emoji     = _DAY_EMOJI[day_key]
            day_title = (entry.get(SynopsisEntryFields.DAY_TITLE, '') if entry else '') or ''
            notes_raw = ''
            if entry:
                polished = entry.get(SynopsisEntryFields.POLISHED_TEXT) or ''
                raw      = entry.get(SynopsisEntryFields.RAW_TEXT) or ''
                notes_raw = polished if polished else (_light_correct(raw) if raw else '')

            p_day = _para(doc, spc_b=12, spc_a=6)
            label = f'{emoji} Day {day_num}'
            if day_title:
                label += f'  ·  {day_title}'
            _run(p_day, label, bold=True, size_pt=14, bg_hex=day_color)

            p_notes = _para(doc, spc_b=6, spc_a=12, line=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            if notes_raw:
                _add_md_text(p_notes, notes_raw, size_pt=11)
            else:
                _run(p_notes, 'No notes submitted for this day.', size_pt=11)

            if entry:
                photo_urls = entry.get(SynopsisEntryFields.PHOTO_URLS, [])
                if photo_urls:
                    for row_start in range(0, len(photo_urls), 3):
                        p_photos = _para(doc, spc_b=4, spc_a=4, align=WD_ALIGN_PARAGRAPH.CENTER)
                        for url in photo_urls[row_start:row_start + 3]:
                            try:
                                img_b = _fix_exif_rotation(_fetch_photo_bytes(url))
                                p_photos.add_run().add_picture(io.BytesIO(img_b), width=Inches(2.1))
                            except Exception as exc:
                                logger.warning(f"Could not embed photo {url}: {exc}")

            p_sep = _para(doc, spc_b=2, spc_a=2)
            _para_border_bottom(p_sep, 'e5e5e5', sz='2')

        _para(doc, spc_b=20, spc_a=20)

    # Food section
    p_brk = doc.add_paragraph()
    p_brk.add_run().add_break(WD_BREAK.PAGE)

    p_fhdr = _para(doc, spc_b=4, spc_a=8)
    _run(p_fhdr, '🍽️  This Week\'s Menu', bold=True, size_pt=15)
    _para_border_bottom(p_fhdr, '1a1a1a', sz='6')

    USABLE_W  = 6.5
    LABEL_COL = 1.3
    DAY_COL   = (USABLE_W - LABEL_COL) / 5

    ftbl = doc.add_table(rows=4, cols=6)
    ftbl.style   = 'Table Grid'
    ftbl.autofit = False
    ftbl.columns[0].width = Inches(LABEL_COL)
    for i in range(1, 6):
        ftbl.columns[i].width = Inches(DAY_COL)

    _cell_bg(ftbl.cell(0, 0), 'f5f5f3')
    for col_i, dk in enumerate(DAY_ORDER):
        cell = ftbl.cell(0, col_i + 1)
        _cell_bg(cell, _DAY_C_HEX[dk])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(_DAY_LABEL_SHORT[dk])
        r.bold      = True
        r.font.size = Pt(12)

    for row_i, (meal_key, meal_label, meal_emoji) in enumerate([
        ('morning_snack',   'Morning Snack',   '🍎'),
        ('lunch',           'Lunch',           '🥗'),
        ('afternoon_snack', 'Afternoon Snack', '🍪'),
    ]):
        tbl_row  = ftbl.rows[row_i + 1]
        lbl_cell = tbl_row.cells[0]
        _cell_bg(lbl_cell, 'f5f5f3')
        p = lbl_cell.paragraphs[0]
        r = p.add_run(f'{meal_emoji} {meal_label}')
        r.bold           = True
        r.font.size      = Pt(11)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

        for col_i, dk in enumerate(DAY_ORDER):
            val  = (food_data.get(dk) or {}).get(meal_key, '') or '—'
            cell = tbl_row.cells[col_i + 1]
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size      = Pt(11)
            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _week_header_meta(week: dict) -> str:
    week_label = week.get(SynopsisWeekFields.LABEL, 'ICC Summer Camps')

    def _ordinal(n: int) -> str:
        if 11 <= n % 100 <= 13:
            return f'{n}th'
        return f'{n}{["th","st","nd","rd","th"][min(n % 10, 4)]}'

    def _fmt_date(s: str) -> str:
        try:
            d = _dt.datetime.strptime(s, '%Y-%m-%d')
            return f'{_ordinal(d.day)} {d.strftime("%B")}'
        except Exception:
            return s

    start = week.get(SynopsisWeekFields.START_DATE, '')
    end   = week.get(SynopsisWeekFields.END_DATE, '')
    if start and end:
        end_year = _dt.datetime.strptime(end, '%Y-%m-%d').year
        dates    = f'{_fmt_date(start)} – {_fmt_date(end)} {end_year}'
    elif start:
        dates = _fmt_date(start)
    else:
        dates = ''
    return week_label + (f' · {dates}' if dates else '')


@router.get("/weeks/{week_id}/download")
async def download_weekly_doc(
    week_id: str,
    group_name: Optional[str] = Query(None, description="Download only this camp group"),
    admin: dict = Depends(verify_icc_admin),
):
    # ── Fetch shared data ──────────────────────────────────────────────────────
    weeks = await firebase.get_all_synopsis_weeks()
    week  = next((w for w in weeks if w.get(SynopsisWeekFields.WEEK_ID) == week_id or w.get('id') == week_id), None)
    if not week:
        raise HTTPException(404, "Week not found")

    all_camps   = await firebase.get_synopsis_camps_for_week(week_id)
    all_entries = await firebase.get_synopsis_entries_for_week(week_id)
    entries_by_camp: dict = {}
    for e in all_entries:
        cid = e.get(SynopsisEntryFields.CAMP_ID)
        entries_by_camp.setdefault(cid, {})[e.get(SynopsisEntryFields.DAY)] = e

    food_data   = await firebase.get_synopsis_food(week_id) or {}
    drive_link  = week.get(SynopsisWeekFields.DRIVE_LINK, '') or ''
    header_meta = _week_header_meta(week)
    year        = _dt.datetime.now().year

    # ── Single-group download → .docx ─────────────────────────────────────────
    if group_name:
        camps = [c for c in all_camps if c.get(SynopsisCampFields.GROUP_NAME) == group_name]
        doc_bytes = _build_synopsis_doc(
            group_camps=camps,
            doc_title=group_name,
            header_meta=header_meta,
            drive_link=drive_link,
            entries_by_camp=entries_by_camp,
            food_data=food_data,
            year=year,
        )
        slug     = group_name.replace(' ', '_')
        filename = f'synopsis_{slug}.docx'
        return StreamingResponse(
            io.BytesIO(doc_bytes),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'},
        )

    # ── All-groups download → .zip of per-group .docx files ───────────────────
    unique_groups = list(dict.fromkeys(
        c.get(SynopsisCampFields.GROUP_NAME, '') for c in all_camps
        if c.get(SynopsisCampFields.GROUP_NAME)
    ))

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for gname in unique_groups:
            camps = [c for c in all_camps if c.get(SynopsisCampFields.GROUP_NAME) == gname]
            doc_bytes = _build_synopsis_doc(
                group_camps=camps,
                doc_title=gname,
                header_meta=header_meta,
                drive_link=drive_link,
                entries_by_camp=entries_by_camp,
                food_data=food_data,
                year=year,
            )
            safe_name = re.sub(r'[^\w\s\-]', '', gname).strip().replace(' ', '_')
            zf.writestr(f'{safe_name}.docx', doc_bytes)

    zip_buf.seek(0)
    return StreamingResponse(
        zip_buf,
        media_type='application/zip',
        headers={'Content-Disposition': f'attachment; filename="synopsis_{week_id}.zip"'},
    )


